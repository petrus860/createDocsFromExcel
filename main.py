"""
-----------------------------------------------------------------------------------------------------------------------
Requisitos para que funcione este script
instalar Python 3.9.5

instalar librerias

		pip install openpyxl
		pip install python-docx

 solucion, error --> raise ValueError('Max value is {0}'.format(self.max)) ValueError: Max value is 52
he buscado en los xml el valor 66 y le he puesto 50

https://stackoverflow.com/questions/50236928/openpyxl-valueerror-max-value-is-14-when-using-load-workbook
Ahora mismo estoy usando el fichero "AT_SP25Mod.xlsx" que funciona correctamente.
Si se quiere reproducit el error 66, hay que usar el nombre del fichero  "AT_SP25.xlsx"

El script funciona de la siguiente manera:
        Leer las columnas ABCDFGH y meter los datos en una lista

        Recorrer lista de cada columna y preparar los datos para crear un archivo doc por cada caso de prueba con sus
            correspondientes pasos de prueba.

        Con los datos crear un archivo en formato .doc con un formato especifico QA

Para mejorar el script --> seria bueno añadir una columna que especifique el tipo de app es CF, ASC, GRA, IBOR
	o preguntar antes por consola
---------------------------------------------------------------------------------------------------------
"""
import openpyxl
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import re
from datetime import date

from operator import itemgetter
from itertools import groupby

xlsx_file = Path('excel', 'Sample_Model0001.xlsx')
wb_obj = openpyxl.load_workbook(xlsx_file)

sheets = wb_obj.sheetnames
# sheet = wb_obj.active
sheet = wb_obj[sheets[0]]   #elegir la hoja 1 de la lista de hojass

print("El nombre actual de la hoja es: ", sheet)

def main():

    contador = 0
    lista_casos_de_prueba = []
    """lee la primera columna A de la hoja excel"""
    for cell in sheet['A']:
        try:
            # print(cell.value)
            id_us = re.findall("BDPDIB-\d{4}", cell.value)
            # print(id_us)
            # if "CIBDP" in cell.value:
            if id_us:
                # print(cell.value)
                valor_columna_a = cell.value
                # valor_columna_a.count("CIBDP")
                contador = contador + 1
                # print(contador)
                lista_casos_de_prueba.append(id_us)

                lista_numeros_columna_f = []
                num_de_pasos = 1

                lista_columna_b = []
                """ lee la columna B de la hoja excel"""
                for cell in sheet['B']:
                    if cell.value != None and cell.value != "US NAME":
                        lista_columna_b.append(cell.value)
                        # print (cell.value)

                lista_columna_c = []
                for cell in sheet['C']: # lee la columna C de la hoja excel
                    try:
                        if cell.value != None and cell.value != "TC ID":
                            lista_columna_c.append(cell.value)
                            # print (cell.value)
                    except:
                        pass

                lista_columna_d = []
                for cell in sheet['D']:
                    try:
                        if cell.value != None and cell.value != "TEST CASE NAME":
                            lista_columna_d.append(cell.value)
                            # print (cell.value)
                    except:
                        pass

                """ lee la columna F de la hoja excel """
                lista_numeros_columna_f = []
                for cell in sheet['F']:  # lee la columna F de la hoja excel
                    try:
                        if type(cell.value) == int:  # me aseguro que tomo solo numeros
                            # print(cell.value)
                            if cell.value >= num_de_pasos:
                                # print("imprimimos el num_de_pasos", num_de_pasos)
                                num_de_pasos = cell.value
                                # print("fin del if cell.value > num_de_pasos")
                                if cell.value == 1:
                                    # print("el numero total de casos de la primera prueba es: ", num_de_pasos)
                                    num_de_pasos = 1

                            if cell.value == 1 and num_de_pasos > cell.value:
                                # print("entro en cell.value menor", num_de_pasos)
                                # print("el numero total de casos de la primera prueba es: ", num_de_pasos)
                                lista_numeros_columna_f.append(num_de_pasos)
                                num_de_pasos = 0
                    except:
                        pass
                # print("el numero total de casos de la primera prueba es: ", num_de_pasos)
                lista_numeros_columna_f.append(num_de_pasos)

                """ lee la columna G de la hoja excel"""
                lista_columna_g = []
                for cell in sheet['G']:
                    try:
                        if cell.value and cell.value != "ACTIONS":
                            lista_columna_g.append(cell.value)
                            # print (cell.value)
                    except:
                        pass
                # print("len lista columna g", len(lista_columna_g))

                """ lee la columna H de la hoja excel """
                lista_columna_h = []
                for cell in sheet['H']:
                    try:
                        if cell.value and cell.value != "EXPECTED RESULT":
                            lista_columna_h.append(cell.value)
                    except:
                        pass
                # print("len lista columna h", len(lista_columna_h))
                # print("valor lista columna h ", lista_columna_h)

                indice_columnas = contador - 1

                """ uso la columna f para que imprima por cada caso cuantos pasos corresponde"""
                indice_columnas_g_h = lista_numeros_columna_f[indice_columnas]

                # print("imprimir lista columna f",lista_numeros_columna_f)

                iteracion_columna_g = iter(lista_columna_g)
                formato_grupo_de_columna_g = [[next(iteracion_columna_g) for _ in range(size)]
                                              for size in lista_numeros_columna_f]
                # print("valor formato_grupo_de_columna_g", formato_grupo_de_columna_g)
                # print("valor lista numeros columna f", lista_numeros_columna_f)

                iteracion_columna_h = iter(lista_columna_h)
                formato_grupo_de_columna_h = [[next(iteracion_columna_h) for _ in range(size)]
                                              for size in lista_numeros_columna_f]
                # print(formato_grupo_de_columna_h)
                # print("valor columna h ", lista_columna_h)


                # print("indice columna F es ", calcular_indice_columna_f)
                # print("Vamos a tratar el caso", contador, "de la HU: ", valor_columna_a)
                print("-------------------------------------------------------------------------")
                # print("valor de indice_columnas_g_h", indice_columnas_g_h)

                crear_archivos_en_formato_word(contador, lista_columna_b[indice_columnas],
                                               lista_columna_c[indice_columnas],
                                               lista_columna_d[indice_columnas],
                                               valor_columna_a, indice_columnas_g_h,
                                               formato_grupo_de_columna_g[indice_columnas],
                                               formato_grupo_de_columna_h[indice_columnas]
                                               )

                crear_fichero_txt(contador, lista_columna_b[indice_columnas],
                                  lista_columna_c[indice_columnas],
                                  lista_columna_d[indice_columnas],
                                  valor_columna_a, indice_columnas_g_h,
                                  formato_grupo_de_columna_g[indice_columnas],
                                  formato_grupo_de_columna_h[indice_columnas]
                                  )
                # print("Vamos a tratar el caso", contador,
                #
                #       "\n Columna B tiene titulo: ", lista_columna_b[indice_columnas],
                #       "\n Columna C tiene titulo: ", lista_columna_c[indice_columnas],
                #       "\n Columna D tiene titulo: ", lista_columna_d[indice_columnas],
                #       "\n de la HU:", valor_columna_a, "tiene ", indice_columnas_g_h, "casos de prueba",
                #       "\n Columna G con sus pasos de prueba: ", formato_grupo_de_columna_g[indice_columnas],
                #       "\n Columna H con sus pasos de prueba: ", formato_grupo_de_columna_h[indice_columnas]
                #       )

                # print("Numero de filas totales es de ", sum(indice_columnas))
                # print("Numero de filas + 1 totales es de ", sum(lista_numeros_columna_f) + 1)
        except:
            pass

        # print(lista_numeros_columna_f, "por el total de ", len(lista_numeros_columna_f), "casos de prueba")

def crear_archivos_en_formato_word(contador, lst_columna_b, lst_columna_c, lst_columna_d,
                                   v_columna_a, indice_columnas_g_h, f_grupo_col_g, f_grupo_col_h):
    print("Ejecuto desde crear_archivos_en_formato_word")
    print("Vamos a tratar el caso", contador,

          "\n Columna B tiene titulo: ", lst_columna_b,
          "\n Columna C tiene titulo: ", lst_columna_c,
          "\n Columna D tiene titulo: ", lst_columna_d,
          "\n de la HU:", v_columna_a, "tiene ", indice_columnas_g_h, "casos de prueba",
          "\n Columna G con sus pasos de prueba: ", f_grupo_col_g,
          "\n Columna H con sus pasos de prueba: ", f_grupo_col_h
          )

    document = Document()
    style = document.styles['Normal']
    fuente = style.font

    fuente.name = 'Arial'
    fuente.size = Pt(11)

    document.add_picture('.\images\Icon1.png', width=Inches(1.75))
    document.add_heading('QA - SCIB', 0) # el 0 aqui significa el estilo
    # parafo = document.add_paragraph()

    fecha_de_hoy = date.today()
    table = document.add_table(rows=1, cols=2)

    table.style = 'Table Grid'

    # table.style = 'Colorful List'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'FECHA'
    hdr_cells[1].text = str(fecha_de_hoy) #'STR0' #poner fecha de hoy
    # hdr_cells[1].text = 'STR0' #poner fecha de hoy
    # hdr_cells[2].text = 'STR00'

    row_cells = table.add_row().cells
    row_cells[0].text = 'APLICACIÓN'
    row_cells[1].text = ''  # poner nombre de la aplicacion
    # row_cells[2].text = 'Str2'

    row2_cells = table.add_row().cells
    row2_cells[0].text = 'RESULTADO DE LA PRUEBA'
    row2_cells[1].text = '' #'Str11'  # poner el resultado de la prueba OK o KO
    # row2_cells[2].text = 'Str22'

    p_vacio = document.add_paragraph('') # hacer una espacio
    p_vacio = document.add_paragraph('') # hacer una espacio

    # document.add_heading('',2)

    p1 = document.add_paragraph(style='List Number').add_run('Definición de la Prueba')
    p1.font.name = 'Arial'
    p1.font.size = Pt(13)
    p1.bold = True
    p1.italic = True

    p1a = document.add_paragraph().add_run(lst_columna_d)
    p1a.font.name = 'Arial'
    p1a.font.size = Pt(10)

    p2 = document.add_paragraph(style='List Number').add_run('Resultados esperados',)
    p2.font.name = 'Arial'
    p2.font.size = Pt(13)
    p2.bold = True
    p2.italic = True

    for i in range(indice_columnas_g_h):
        p2a = document.add_paragraph(style='List Bullet').add_run(f_grupo_col_h[i])  # resultados esperado
        p2a.font.name = 'Arial'
        p2a.font.size = Pt(10)

    p3 = document.add_paragraph(style='List Number').add_run('Evidencias funcionales',)
    p3.font.name = 'Arial'
    p3.font.size = Pt(13)
    p3.bold = True
    p3.italic = True

    for i in range(indice_columnas_g_h):
        p3a = document.add_paragraph(style='List Number 2').add_run(f_grupo_col_g[i]) # acciones
        p3a.font.name = 'Arial'
        p3a.font.size = Pt(10)

        # for i_h in len(f_grupo_col_h):
        p3b = document.add_paragraph(style='List Bullet 3').add_run(f_grupo_col_h[i]) # resultados esperado
        # p3b = document.add_paragraph(style='List Number 3').add_run(f_grupo_col_h[i]) # resultados esperado
        p3b.font.name = 'Arial'
        p3b.font.size = Pt(10)

    # document.add_page_break()

    nombre_fichero_word = "00" + str(contador) + "_" + "HU" + "_" + str(v_columna_a) + ".docx"
    print("valor nombre fichero: ", nombre_fichero_word)

    document.save('.\\doc\\' + nombre_fichero_word)



def crear_fichero_txt(contador, lst_columna_b, lst_columna_c, lst_columna_d,
                                   v_columna_a, indice_columnas_g_h, f_grupo_col_g, f_grupo_col_h):

    nombre_fichero_txt = "00" + str(contador) + "_" + "HU" + "_" + str(v_columna_a) + ".txt"
    print("valor nombre fichero: ", nombre_fichero_txt)
    fichero_txt = open(".\\txt\\" + nombre_fichero_txt, "w")


    # fichero_txt.write("esto es un texxto")
    fichero_txt.write('{}\n{}\n{}\n{}\n'.format(contador, lst_columna_b, lst_columna_c,lst_columna_d))
    fichero_txt.write("\n")
    for i in f_grupo_col_g:
        fichero_txt.write(i)
        fichero_txt.write("\n")
    fichero_txt.write("\n")
    fichero_txt.write("\n")
    for i in f_grupo_col_h:
        # fichero_txt.writelines('{}\n{}\n'.format(f_grupo_col_g, f_grupo_col_h))
        fichero_txt.write(i)
        fichero_txt.write("\n")

    fichero_txt.close()


main()
# primer_Caso_mas_doc()
# crear_diccionario_de_todo_excel()
# primer_Caso_mas_doc()
